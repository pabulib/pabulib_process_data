import collections
import os
import re
from dataclasses import dataclass

import helpers.utilities as utils
import openpyxl
import pypdf
from process_data.cities.gdynia.postprocess import register_gdynia_postprocess
from process_data.base_config import BaseConfig
from process_data.models import ProjectItem
import xlrd


@dataclass(kw_only=True)
class GetProjects(BaseConfig):
    data_dir: str
    projects_docx: str = None
    columns: dict = None
    cell_colours: dict = None
    subdistricts: bool = True
    city_projects_filename: str = None
    district_projects_filename: str = None
    results_pdf: str = None

    def __post_init__(self):
        self.logger = utils.create_logger()
        self.initialize_mapping_dicts()
        if self.instance >= 2021:
            self.results_mapping = self.get_results_mapping_from_pdf()
        else:
            self.check_if_pdf()
        return super().__post_init__()

    def get_results_mapping_from_pdf(self):
        path_to_pdf = utils.get_path_to_file_by_unit(
            self.results_pdf, self.unit, self.data_dir, "pdf"
        )
        reader = pypdf.PdfReader(path_to_pdf)
        text = "\n".join(page.extract_text() or "" for page in reader.pages)
        matches = re.finditer(
            r"((?:20\d{2}/)?[A-Z]{3}/\d{4}).*?(\d[\d ]*)\s*(?:zł)?\s+(\d+)\s+(TAK|NIE|tak|nie|PROJEKT \+1|Projekt \+1|projekt \+1)",
            text,
            re.DOTALL,
        )
        results_mapping = {}
        for match in matches:
            project_id, _, votes, selected_txt = match.groups()
            if selected_txt.lower() == "tak":
                selected = 1
            elif selected_txt.lower() == "projekt +1":
                selected = 2
            else:
                selected = 0
            results_mapping[project_id] = {"votes": int(votes), "selected": selected}
        return results_mapping

    def check_if_pdf(self):
        import pdf2docx

        if self.projects_docx.endswith(".pdf"):
            pdf_name = self.projects_docx.replace(".pdf", "")
            path_to_pdf = utils.get_path_to_file_by_unit(
                pdf_name, self.unit, self.data_dir, "pdf"
            )

            self.path_to_docx = utils.get_path_to_file_by_unit(
                pdf_name, self.unit, self.data_dir, "docx"
            )

            pdf2docx.parse(path_to_pdf, self.path_to_docx)

        else:
            self.path_to_docx = utils.get_path_to_file_by_unit(
                self.projects_docx, self.unit, self.data_dir, "docx"
            )

    def get_cell_colour(self, cell):
        pattern = re.compile('w:fill="(\S*)"')
        match = pattern.search(cell._tc.xml)
        try:
            result = match.group(1)
        except AttributeError:
            result = None
        return result

    def check_if_bold_text(self, cell):
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                if run.bold:
                    return True
                return

    def check_project_id(self, project_id):
        project_id = project_id.strip()
        if "\t" in project_id:
            project_id = project_id.split("\t")[1]
        return project_id

    def get_district_subdistrict_from_merged_cell(self, value):
        if "miejskie ogólne" in value:
            district = "CITYWIDE"
            subdistrict = "CITYWIDE"
        elif "miejske klimatyczny" in value:
            district = "GREEN"
            subdistrict = "GREEN"
        elif "projekty duże" in value:
            subdistrict = "large"
            value = value.replace(" projekty duże", "")
            district = value.title()
        elif "projekty małe" in value:
            subdistrict = "small"
            value = value.replace(" projekty małe", "")
            district = value.title()
        else:
            raise RuntimeError(f"Different value in merged cell! {value}")
        return district, subdistrict

    def get_data_from_docx_table(self, doc):
        for page_no in range(len(doc.tables)):
            table = doc.tables[page_no]
            for i, row in enumerate(table.rows):
                row_values = [cell.text for cell in row.cells]

                # Establish the mapping based on the first row
                # headers; these will become the keys of our dictionary
                if page_no == 0:
                    if i == 0:
                        continue
                    elif i == 1:
                        values = [cell.text.replace("\n", "")
                                  for cell in row.cells]
                        keys = [self.columns.get(value, value)
                                for value in values]
                        continue

                unique_values = set(row_values)
                cell_colour = self.get_cell_colour(row.cells[0])

                if len(unique_values) == 1:
                    value = list(unique_values)[0]
                    if self.instance == 2020:
                        if value == "Ogólnomiejskie":
                            district = "CITYWIDE"
                            subdistrict = "CITYWIDE"
                            continue

                        if cell_colour.lower() == self.cell_colours["district"]:
                            district = value.title()
                            continue
                        elif cell_colour.lower() == self.cell_colours["size"]:
                            if "małe" in value.lower():
                                subdistrict = "small"
                            elif "duże" in value.lower():
                                subdistrict = "large"
                            continue
                    elif self.instance > 2020:
                        district, subdistrict =\
                            self.get_district_subdistrict_from_merged_cell(
                                value.lower())
                        continue

                # Construct a dictionary for this row, mapping keys to values
                row_data = dict(zip(keys, row_values))
                # print(row.cells[0]._tc.xml.w)

                selected = self.check_if_selected_project(row_data)
                row_data["selected"] = selected
                row_data["cost"] = utils.get_cost_from_text(row_data["cost"])
                row_data["votes"] = int(row_data["votes"].replace(" ", ""))
                row_data["name"] = row_data["name"].replace("\n", "")
                row_data["project_id"] = self.check_project_id(
                    row_data["project_id"])

                row_data = {k: row_data[k] for k in self.columns.values()}

                project_id = row_data["project_id"]
                if self.subdistricts:
                    if not self.projects_data_per_district.get(district):
                        self.projects_data_per_district[
                            district
                        ] = collections.defaultdict(list)
                        self.district_projects_mapping[
                            district
                        ] = collections.defaultdict(list)

                    self.projects_data_per_district[district][subdistrict].append(
                        row_data
                    )
                    self.district_projects_mapping[district][subdistrict].append(
                        project_id
                    )
                    self.project_district_mapping[project_id] = district
                    self.project_subdistrict_mapping[project_id] = subdistrict

    def check_if_selected_project(self, row_data):
        selected_text = row_data["selected"].replace("\n", "")
        if selected_text == "tak":
            return 1
        if (selected_text == "Projekt +1") or ("Asseco" in selected_text):
            # "sfinansowany przez Asseco Data Systems"
            # Additional funded project
            return 2
        return 0

    def initialize_mapping_dicts(self):
        self.projects_data_per_district = dict()
        self.district_projects_mapping = collections.defaultdict(list)
        self.project_district_mapping = dict()
        self.project_subdistrict_mapping = dict()
        self.project_citywide_pool_mapping = dict()

    def create_district_upper_mapping(self):
        self.district_upper_district_mapping = {
            "CITYWIDE": "citywide"
        }
        for district in self.district_projects_mapping.keys():
            district_upper = utils.change_district_into_name(district)
            self.district_upper_district_mapping[district_upper] = district

    def create_project_item(self):
        project_id = None
        item = ProjectItem(project_id)
        name = None
        item.add_name(name)
        district = None
        self.logger.info(f"Processing Project: {project_id}")
        item.add_district(district.strip())
        item.project_url = None
        cost = None
        item.add_cost(cost)
        status = None
        item.add_selected(status)
        item.votes = None

        return item

    def get_workbook_rows(self, filename):
        csv_path = utils.get_path_to_file_by_unit(filename, self.unit, self.data_dir, "csv")
        xls_path = utils.get_path_to_file_by_unit(filename, self.unit, self.data_dir, "xls")
        xlsx_path = utils.get_path_to_file_by_unit(
            filename, self.unit, self.data_dir, "xlsx"
        )
        if os.path.exists(csv_path):
            import csv

            with open(csv_path, encoding="utf-8-sig", newline="") as csv_file:
                reader = csv.DictReader(csv_file, delimiter=";")
                return list(reader)
        if os.path.exists(xls_path):
            workbook = xlrd.open_workbook(xls_path)
            sheet = workbook.sheet_by_index(0)
            header = [sheet.cell_value(0, col) for col in range(sheet.ncols)]
            return [
                {header[col]: sheet.cell_value(row_no, col) for col in range(sheet.ncols)}
                for row_no in range(1, sheet.nrows)
            ]

        workbook = openpyxl.load_workbook(xlsx_path, read_only=True)
        sheet = workbook[workbook.sheetnames[0]]
        rows_iter = sheet.iter_rows(values_only=True)
        header = list(next(rows_iter))
        return [dict(zip(header, row_values)) for row_values in rows_iter]

    def get_first_present_value(self, row_data, *keys):
        for key in keys:
            if key in row_data and row_data[key] not in (None, ""):
                return row_data[key]
        return None

    def get_citywide_pool(self, row_data):
        task_type = str(
            self.get_first_present_value(row_data, "Typ zadania", "typ_projektu") or ""
        ).strip()
        if "Klimatycznego Budżetu Obywatelskiego" in task_type:
            return "KBO"
        if "Jubileuszowego Budżetu Obywatelskiego" in task_type:
            return "JBO"
        if task_type == "KBO":
            return "KBO"
        if task_type in ("Ogólny", "") and self.instance == 2021:
            return "GENERAL"
        if self.instance in (2022, 2023):
            return "KBO"
        return None

    def get_subdistrict_2021_plus(self, row_data):
        district_value = self.get_first_present_value(row_data, "Dzielnica", "dzielnica")
        district = str(district_value or "").strip()
        if not district or district.lower().startswith("ogólnomi"):
            return "CITYWIDE", "CITYWIDE"

        project_type = str(
            self.get_first_present_value(row_data, "Typ projektu", "typ_projektu") or ""
        ).strip().lower()
        if project_type == "mały":
            return district, "small"
        if project_type == "duży":
            return district, "large"
        raise RuntimeError(f"Unknown Gdynia project type: {project_type}")

    def get_project_rows_from_xls(self, filename):
        for row_data in self.get_workbook_rows(filename):
            project_id = str(
                self.get_first_present_value(row_data, "ID projektu", "id_projektu", "id_wniosku")
            ).strip()
            if project_id not in self.results_mapping:
                continue
            yield row_data

    def get_data_from_2021_plus_files(self):
        for filename in (
            self.city_projects_filename,
            self.district_projects_filename,
        ):
            for row_data in self.get_project_rows_from_xls(filename):
                district, subdistrict = self.get_subdistrict_2021_plus(row_data)
                project_id = str(
                    self.get_first_present_value(row_data, "ID projektu", "id_projektu", "id_wniosku")
                ).strip()
                result = self.results_mapping[project_id]
                row = {
                    "project_id": project_id,
                    "name": str(
                        self.get_first_present_value(row_data, "Tytuł", "Tytuł projektu", "tytul_projektu")
                    ).replace("\n", "").strip(),
                    "cost": utils.get_cost_from_text(
                        str(self.get_first_present_value(row_data, "Szacunkowy koszt", "szacunkowy_koszt"))
                    ),
                    "votes": result["votes"],
                    "selected": result["selected"],
                }

                if not self.projects_data_per_district.get(district):
                    self.projects_data_per_district[district] = collections.defaultdict(
                        list
                    )
                    self.district_projects_mapping[district] = collections.defaultdict(
                        list
                    )

                self.projects_data_per_district[district][subdistrict].append(row)
                self.district_projects_mapping[district][subdistrict].append(project_id)
                self.project_district_mapping[project_id] = district
                self.project_subdistrict_mapping[project_id] = subdistrict
                if district == "CITYWIDE":
                    pool = self.get_citywide_pool(row_data)
                    if pool:
                        self.project_citywide_pool_mapping[project_id] = pool

    def start(self):
        if self.instance >= 2020:
            register_gdynia_postprocess(
                country=self.country,
                unit=self.unit,
                instance=int(self.instance),
            )

        if self.instance >= 2021:
            self.get_data_from_2021_plus_files()
            self.create_district_upper_mapping()
            objects = {
                "district_projects_mapping": self.district_projects_mapping,
                "projects_data_per_district": self.projects_data_per_district,
                "project_district_mapping": self.project_district_mapping,
                "project_subdistrict_mapping": self.project_subdistrict_mapping,
                "project_citywide_pool_mapping": self.project_citywide_pool_mapping,
                "district_upper_district_mapping": self.district_upper_district_mapping,
            }
            self.save_mappings_as_jsons(objects)
            return

        import docx

        docx_file = docx.Document(self.path_to_docx)
        self.get_data_from_docx_table(docx_file)
        self.create_district_upper_mapping()
        # self.create_district_name_mapping()
        objects = {
            "district_projects_mapping": self.district_projects_mapping,
            "projects_data_per_district": self.projects_data_per_district,
            "project_district_mapping": self.project_district_mapping,
            "project_subdistrict_mapping": self.project_subdistrict_mapping,
            "district_upper_district_mapping": self.district_upper_district_mapping,
        }
        self.save_mappings_as_jsons(objects)
