import os
import docx

pabulib_dir = os.path.join(os.getcwd(), "src", "pabulib")

path_to_pdf_files = os.path.join(
    pabulib_dir, "data", "Czestochowa", "edycja_6_2020")
output_doc_name = "sample.docx"

path_to_doc = os.path.join(
    path_to_pdf_files, output_doc_name).replace("\\", "/")

doc = docx.Document(path_to_doc)


def get_data_from_docx_table(doc):
    district_data = list()
    keys = None
    for page_no in range(len(doc.tables)):
        table = doc.tables[page_no]

        for i, row in enumerate(table.rows):
            text = None
            text = (cell.text for cell in row.cells)

            # Establish the mapping based on the first row
            # headers; these will become the keys of our dictionary
            if page_no == 0:
                if i == 0:
                    continue
                if i == 1:
                    keys = tuple(text)
                    continue

            # Construct a dictionary for this row, mapping
            # keys to values for this row
            row_data = dict(zip(keys, text))
            district_data.append(row_data)
    return district_data


district_data = get_data_from_docx_table(doc)
print(district_data)
